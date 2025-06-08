#все нужно установить в терминале
#pip install python-docx
#pip install pandas
#pip install openpyxl
#pip install docxtpl
#pip install spire.Doc
#pip install os


#Расшифровка

# Пути к файлам, в конце предоставлено, как должно выглядеть окончание ссылки

#ссылка на список жильцов в Excel (берется отсуп в 5 ячеек)
#Excel_scr = "C:/_/_/_/_/готовый список.xlsx"

#представители  в Excel (список на представителей вместо жильцов) (важно присутствие в таблице точного ФИО жильца,
                                                                    #в проттивном случае - заполнения не будет)
#Excel_agent_scr = "C:/_/_/_/_/Представитель собственника.xlsx" 

#шаблон Word (размеченный)
#Word_src = "C:/_/_/_/_/шаблон.docx" 

#ссылка на конечный путь сохранения (единого файла), создается файл, в котором будут склеены все файлы (для распечатки)
#out_src = "C:/_/_/_/_/наполнитель.docx"      

#ссылка на конечный путь сохранения (на каждую персону отдельный файл), создается папка, в которой будут вложены все файлы
#temp_folder = "C:/_/_/_/_/заполненные"       

#вопросы для голосования
#questions_src = "C:/_/_/_/_/вопросы.xlsx"  




from openpyxl import load_workbook
from docxtpl import DocxTemplate
from spire.doc import Document as SpireDocument, FileFormat
import os
from collections import defaultdict
from docx import Document


# Создаем папку для временных файлов
os.makedirs(temp_folder, exist_ok=True)

# Открываем файлы Excel
wb = load_workbook(Excel_scr)
sheet = wb.active

wb_agent = load_workbook(Excel_agent_scr)
sheet_agent = wb_agent.active

wb_questions = load_workbook(questions_src)
sheet_questions = wb_questions.active

# Итоговый документ
final_doc = SpireDocument()

# Статические данные
home_data = {
    "home_number": "100",
    "home_street": "ул. Красная",
    "home_town": "г. Краснодар",
    "home_sqm": "4200",
    "home_votes": "100"
}

# Дата предоставления протоколов
date_time = {
    "date_votes": "17.03.2025"  
}

# Загружаем вопросы из Excel
questions_list = []
for row in sheet_questions.iter_rows(min_row=2, values_only=True):  # Пропускаем заголовок
    question_number = row[0]  # Столбец A - номер вопроса
    question_text = row[1]  # Столбец B - сам вопрос
    if question_number and question_text:  # Пропускаем пустые строки
        questions_list.append({"question_number": question_number, "question_text": question_text})

# Создаём пустой словарь, если его ещё нет
union_dictionary = {}

# Добавляем список вопросов
union_dictionary["questions"] = questions_list

# Фильтрация строк с жильцами
rows_with_data = [
    row for row in sheet.iter_rows(min_row=5)
    if any(cell.value is not None for cell in row)
]

# Словарь для быстрого поиска представителя
agent_dict = {}
for agent_row in sheet_agent.iter_rows(min_row=2, values_only=True):  # Пропускаем заголовок
    agent_name = agent_row[2]  # Столбец C (индекс 2)
    agent_dict[agent_name] = {
        "owners_agent_name": agent_row[1],  # Столбец B
        "number_owners_agent": agent_row[5],  # Столбец F
        "owners_agent_date": agent_row[4]  # Столбец E
    }

# Счетчик жильцов по квартирам
apartment_resident_counter = defaultdict(int)

for row in rows_with_data:
    # Данные жильца
    resident_data = {
        "resident_name": row[1].value,  # Столбец B
        "resident_apartment": row[2].value,  # Столбец C
        "resident_ownership_basis": row[4].value,  # Столбец E
        "resident_sqm": row[6].value,  # Столбец G
        "resident_number_vote": row[8].value  # Столбец I
    }

    # Увеличиваем счетчик для данной квартиры
    apartment_resident_counter[resident_data["resident_apartment"]] += 1
    resident_position = apartment_resident_counter[resident_data["resident_apartment"]]

    # Поиск представителя
    agent_data = agent_dict.get(resident_data["resident_name"], {  # Если представителя нет, то оставляем пустые значения
        "owners_agent_name": "",
        "number_owners_agent": "",
        "owners_agent_date": ""
    })

    # Объединяем данные
    union_dictionary = {
        **home_data,
        **date_time,
        **resident_data,
        **agent_data,
        "questions": questions_list  # Добавляем список вопросов
    }

    # Создаем копию шаблона и рендерим
    doc_template = DocxTemplate(Word_src)
    doc_template.render(union_dictionary)

    # Получаем номер квартиры
    apartment_number = resident_data["resident_apartment"]

    # Формируем имя файла с учетом порядкового номера жильца
    if resident_position == 1:
        file_name = f"0{apartment_number} квартира.docx"
    else:
        file_name = f"0{apartment_number}_{resident_position} квартира.docx"

    # Путь для временного файла
    temp_doc_path = os.path.join(temp_folder, file_name)
    doc_template.save(temp_doc_path)


    # Функция для удаления пустых строк из таблиц
    def remove_empty_rows(doc_path):
        doc = Document(doc_path)

        for table in doc.tables: 
            rows_to_delete = []
        for i, row in enumerate(table.rows):
            if all(cell.text.strip() == "" for cell in row.cells):
                rows_to_delete.append(i)

        for i in reversed(rows_to_delete):
            table._element.remove(table.rows[i]._element)
            
        doc.save(doc_path)

    # Удаляем пустые строки из таблицы в документе после рендеринга
    remove_empty_rows(temp_doc_path)

    # Вставляем в итоговый документ
    final_doc.InsertTextFromFile(temp_doc_path, FileFormat.Auto)

# Сохранение итогового документа
final_doc.SaveToFile(out_src, FileFormat.Docx)
final_doc.Close()

print(f"Документ успешно сохранен в: {out_src}")
